"""文档解析器：支持 PDF / Word / Markdown / HTML / TXT → 统一文本块列表"""
from __future__ import annotations

import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Generator

logger = logging.getLogger(__name__)


@dataclass
class ParsedBlock:
    """解析出的文本块"""
    content: str
    metadata: dict = field(default_factory=dict)
    # metadata 可包含: page_num, section_title, is_table, is_code, source_type


class DocumentParser:
    """
    统一文档解析入口。
    
    支持格式：
      - PDF  → pdfplumber（保留版面、表格）
      - Word → python-docx（段落 + 表格）
      - Markdown → 直接读取
      - HTML  → html2text / BeautifulSoup
      - TXT   → 直接读取
    
    输出：Generator[ParsedBlock]，每个 block 是一个语义单元（段落/表格/标题）。
    """

    def __init__(self):
        self._parsers = {
            ".pdf": self._parse_pdf,
            ".docx": self._parse_docx,
            ".doc": self._parse_docx,  # .doc 也尝试用 docx 库处理
            ".md": self._parse_markdown,
            ".markdown": self._parse_markdown,
            ".html": self._parse_html,
            ".htm": self._parse_html,
            ".txt": self._parse_text,
        }

    def parse(self, file_path: str | Path) -> Generator[ParsedBlock, None, None]:
        """
        根据文件扩展名自动选择解析器，生成文本块序列。
        
        Args:
            file_path: 文件路径
            
        Yields:
            ParsedBlock: 解析出的文本块
        """
        path = Path(file_path)
        ext = path.suffix.lower()

        parser_func = self._parsers.get(ext)
        if not parser_func:
            logger.warning(f"不支持的文件类型: {ext}，尝试按纯文本读取")
            yield from self._parse_text(path)
            return

        try:
            yield from parser_func(path)
        except Exception as e:
            logger.error(f"解析文件 {file_path} 失败: {e}")
            raise

    # ── 各格式解析实现 ──────────────────────────────

    def _parse_pdf(self, path: Path) -> Generator[ParsedBlock, None, None]:
        """PDF 解析：逐页提取文本 + 表格"""
        import pdfplumber

        with pdfplumber.open(str(path)) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                # 提取页面文本
                text = page.extract_text() or ""
                if text.strip():
                    yield ParsedBlock(
                        content=text.strip(),
                        metadata={"page_num": page_num, "source_type": "pdf"},
                    )

                # 提取表格
                tables = page.extract_tables()
                for table_idx, table in enumerate(tables or []):
                    if not table:
                        continue
                    rows = []
                    for row in table:
                        cell_text = " | ".join(
                            (cell or "").strip().replace("\n", " ") for cell in row
                        )
                        rows.append(cell_text)
                    table_str = "\n".join(rows)
                    if table_str.strip():
                        yield ParsedBlock(
                            content=table_str,
                            metadata={
                                "page_num": page_num,
                                "source_type": "pdf",
                                "is_table": True,
                                "table_index": table_idx,
                            },
                        )

    def _parse_docx(self, path: Path) -> Generator[ParsedBlock, None, None]:
        """Word 文档解析：段落 + 表格"""
        from docx import Document

        doc = Document(str(path))

        # 段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                style_name = para.style.name if para.style else ""
                yield ParsedBlock(
                    content=text,
                    metadata={
                        "source_type": "docx",
                        "style": style_name,
                        "is_heading": "Heading" in style_name or "Title" in style_name,
                    },
                )

        # 表格
        for table_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [(cell.text or "").strip().replace("\n", " ") for cell in row.cells]
                rows.append(" | ".join(cells))
            table_str = "\n".join(rows)
            if table_str.strip():
                yield ParsedBlock(
                    content=table_str,
                    metadata={"source_type": "docx", "is_table": True, "table_index": table_idx},
                )

    def _parse_markdown(self, path: Path) -> Generator[ParsedBlock, None, None]:
        """Markdown 解析：按空行分块，保留标题层级"""
        content = path.read_text(encoding="utf-8")
        blocks = content.split("\n\n")

        for block in blocks:
            text = block.strip()
            if not text:
                continue
            lines = text.split("\n")
            first_line = lines[0] if lines else ""
            is_heading = first_line.startswith("#")
            yield ParsedBlock(
                content=text,
                metadata={
                    "source_type": "markdown",
                    "is_heading": is_heading,
                    "heading_level": first_line.count("#") if is_heading else 0,
                },
            )

    def _parse_html(self, path: Path) -> Generator[ParsedBlock, None, None]:
        """HTML 解析：提取正文内容，去掉标签噪声"""
        from bs4 import BeautifulSoup

        html_content = path.read_text(encoding="utf-8")
        soup = BeautifulSoup(html_content, "lxml")

        # 去掉 script/style
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()

        text = soup.get_text(separator="\n", strip=True)
        for block in text.split("\n\n"):
            block = block.strip()
            if block and len(block) > 10:  # 过滤太短的碎片
                yield ParsedBlock(content=block, metadata={"source_type": "html"})

    def _parse_text(self, path: Path) -> Generator[ParsedBlock, None, None]:
        """纯文本解析：按段落分块"""
        content = path.read_text(encoding="utf-8")
        for block in content.split("\n\n"):
            block = block.strip()
            if block:
                yield ParsedBlock(content=block, metadata={"source_type": "text"})


# 全局单例
parser = DocumentParser()
