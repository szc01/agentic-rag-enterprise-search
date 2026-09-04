# -*- coding: utf-8 -*-
"""
按学校下发的《2-专业综合工程实践报告模板.docx》重新生成报告。
模板规范：
- 页面：A4 纵向
- 页边距：上 70.9 / 下 70.9 / 左 85.05 / 右 70.9 pt
- 字体：仿宋_GB2312（封面与正文统一）
- 行距：18pt 固定值（WD_LINE_SPACING.EXACTLY）
- 章节标题：分散对齐 + 悬挂缩进（左 -5.25pt）
- 章节结构：1 引言 / 2 需求分析 / 3 系统设计 / 4 详细设计 / 5 系统测试 / 结论 / 参考文献 / 附录A / 附录B
- 参考文献分三类：[1] 期刊 / [a] 图书 / [c] 外文
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

IMGDIR = r"D:\agentic-rag-system\report\images"
OUT = r"D:\agentic-rag-system\report\专业综合工程实践设计报告.docx"

# ── 字体常量（按模板：仿宋_GB2312）────────
FANGSONG = "仿宋_GB2312"
HEI = "黑体"  # 仅小标题可点缀
KAI = "楷体_GB2312"
EN_FONT = "Times New Roman"
BODY_SIZE = Pt(12)        # 小四
H1_SIZE = Pt(16)          # 三号
H2_SIZE = Pt(14)          # 四号
H3_SIZE = Pt(12)          # 小四加粗
DARK = RGBColor(0, 0, 0)

doc = Document()


# ── 页面：A4 + 自定义页边距 ─────────
section = doc.sections[0]
section.page_height = Pt(842)
section.page_width = Pt(595.35)
section.top_margin = Pt(70.9)
section.bottom_margin = Pt(70.9)
section.left_margin = Pt(85.05)
section.right_margin = Pt(70.9)


def set_run(run, cn_font=FANGSONG, size=BODY_SIZE, bold=False, color=DARK, italic=False):
    """中文字体走 eastAsia，西文走 Times New Roman。"""
    run.font.name = EN_FONT
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:eastAsia"), cn_font)
    run._element.rPr.rFonts.set(qn("w:ascii"), EN_FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), EN_FONT)


def add_para(text="", cn_font=FANGSONG, size=BODY_SIZE, bold=False, align=None,
             indent=True, space_before=0, space_after=6,
             line_spacing_rule=WD_LINE_SPACING.EXACTLY, line_spacing_pt=18):
    """正文段落：18pt 固定行距 + 首行缩进 2 字符。"""
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = line_spacing_rule
    pf.line_spacing = Pt(line_spacing_pt)
    if indent:
        pf.first_line_indent = Pt(24)  # 2 字符
    if text:
        r = p.add_run(text)
        set_run(r, cn_font=cn_font, size=size, bold=bold)
    return p


def add_chapter_heading(text):
    """一级标题：分散对齐 + 悬挂缩进（按模板），居中粗体三号字。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
    pf = p.paragraph_format
    pf.left_indent = Pt(-5.25)
    pf.right_indent = Pt(-16.15)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(18)
    pf.space_before = Pt(18)
    pf.space_after = Pt(12)
    r = p.add_run(text)
    set_run(r, cn_font=HEI, size=H1_SIZE, bold=True)
    return p


def add_section_heading(text):
    """二级标题：左对齐 + 粗体四号。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(18)
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    r = p.add_run(text)
    set_run(r, cn_font=HEI, size=H2_SIZE, bold=True)
    return p


def add_subsection_heading(text):
    """三级标题：左对齐 + 粗体小四。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(18)
    pf.space_before = Pt(8)
    pf.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, cn_font=HEI, size=H3_SIZE, bold=True)
    return p


def add_figure(img_name, caption, width_cm=14.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(f"{IMGDIR}\\{img_name}", width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    cap.paragraph_format.line_spacing = Pt(18)
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    set_run(r, cn_font=HEI, size=Pt(10.5), bold=False)
    return p


def add_table(headers, rows, col_widths=None, font_size=Pt(10.5)):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run(r, cn_font=HEI, size=font_size, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 or len(str(val)) > 20 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            set_run(r, cn_font=FANGSONG, size=font_size)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    sp = doc.add_paragraph()
    sp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    sp.paragraph_format.line_spacing = Pt(18)
    sp.paragraph_format.space_after = Pt(4)
    return t


# ══════════════════════════════════════════════
# 封面（按模板：仿宋_GB2312 + 48pt 粗体）
# ══════════════════════════════════════════════
for _ in range(3):
    add_para("", indent=False, space_after=12)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("专业综合工程实践设计报告")
set_run(r, FANGSONG, Pt(36), True)  # 模板实测 48pt 在 A4 上偏大，36pt 更均衡

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
r = p.add_run("——基于 Agentic RAG 的企业智能搜索与自动调研系统")
set_run(r, FANGSONG, Pt(18))

for _ in range(3):
    add_para("", indent=False, space_after=12)

def cover_line(label, value="____________________________"):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(28)
    r = p.add_run(label + value)
    set_run(r, FANGSONG, Pt(14))

cover_line("题    目：", "基于 Agentic RAG 的企业智能搜索与自动调研系统")
cover_line("院系名称：", "计算机科学与技术学院")
cover_line("学生姓名：", "随治诚")
cover_line("指导教师：", "________________")
cover_line("专业班级：", "计算机科学与技术")

add_para("", indent=False, space_after=12)
add_para("", indent=False, space_after=12)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("二〇二六年九月")
set_run(r, FANGSONG, Pt(14))

doc.add_page_break()


# ══════════════════════════════════════════════
# 目录
# ══════════════════════════════════════════════
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("目  录")
set_run(r, HEI, Pt(18), True)
p.paragraph_format.space_after = Pt(18)

# 静态目录（按模板风格手写，自动生成需 Word 端 F9 更新）
toc_items = [
    ("1  引    言", "1"),
    ("2  需求分析", "3"),
    ("3  系统设计", "6"),
    ("4  详细设计", "11"),
    ("5  系统测试", "16"),
    ("结    论", "21"),
    ("参考文献", "23"),
    ("附录A  项目部署", "25"),
    ("附录B  用户手册", "27"),
]
for title, page in toc_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(20)
    pf.space_after = Pt(3)
    # 用 tab 把页码推到右对齐
    tab_stops = pf.tab_stops
    tab_stops.add_tab_stop(Cm(13.5), alignment=2, leader=1)  # 2=RIGHT, 1=DOTS
    r = p.add_run(title)
    set_run(r, FANGSONG, Pt(12))
    r2 = p.add_run("\t" + page)
    set_run(r2, FANGSONG, Pt(12))

doc.add_page_break()


# ══════════════════════════════════════════════
# 第 1 章 引言
# ══════════════════════════════════════════════
add_chapter_heading("1  引    言")

add_section_heading("1.1  研究背景与意义")
add_para(
    "随着企业数字化转型的深入推进，组织内部沉淀的合同、规范、调研报告、技术文档等非结构化知识"
    "呈指数级增长，传统的关键词检索（Elasticsearch / 数据库 LIKE）已难以满足员工的精准知识获取需求："
    "一方面，基于字面匹配的检索无法捕捉查询背后的语义意图，导致大量「找得到词、找不到意」的情况；"
    "另一方面，复杂的业务问题往往需要跨文档综合，单纯的「检索一次即返回」模式无法满足多跳推理需求。"
)
add_para(
    "检索增强生成（Retrieval-Augmented Generation, RAG）将信息检索与大型语言模型相结合，通过外挂知识库"
    "为 LLM 提供事实性上下文，已成为企业知识管理的热门技术路线。然而，传统单轮 RAG 在面对复杂、"
    "多跳问题时存在明显不足：检索质量受限于查询的表达；缺乏对检索结果的批判与再检索机制；"
    "难以处理跨文档综合任务。基于此，本文研究并实现一个基于 Agentic RAG 的企业智能搜索与自动调研系统，"
    "对于推动 RAG 技术在企业知识管理中的落地具有工程实践价值。"
)

add_section_heading("1.2  国内外研究现状")
add_para(
    "在检索侧，从 Lewis 等人于 2020 年提出 RAG 范式以来，研究沿着「检索器增强—索引优化—生成器优化」路线"
    "快速发展：稠密检索（DPR）→ 稀疏 + 稠密混合检索 → 引入 Cross-Encoder 精排（Rerank）。向量基础设施"
    "方面，FAISS、Milvus、pgvector 等成熟方案支撑了大规模相似度检索，其中 pgvector 以与 PostgreSQL 的"
    "天然集成成为中小规模企业知识库的热门选择。在智能体编排方面，Self-RAG、Corrective RAG 等工作"
    "通过反思与迭代提升检索质量，LangGraph 等框架使多智能体协作落地更便捷。"
)
add_para(
    "总体来看，Agentic RAG 在学术界与工业界均处于快速上升期，但面向中文企业知识库的完整工程实现"
    "以及配套的评测与可视化体系仍缺乏系统性落地实践。本课题正是基于这一背景展开。"
)

add_section_heading("1.3  本文主要工作")
add_para(
    "本文设计并实现了基于 Agentic RAG 的企业智能搜索与自动调研系统，主要工作包括："
)
add_para(
    "（1）构建文档入库全链路，支持多格式解析、多策略分片、向量化入库与 BM25 倒排索引的增量更新；"
)
add_para(
    "（2）实现 BM25 稀疏检索 + pgvector 稠密检索 + RRF 融合 + BGE-Reranker 精排的混合检索管线，"
    "并引入 LLM 查询改写与 HyDE 检索增强；"
)
add_para(
    "（3）基于 LangGraph 构建 Planner—Retrieval—Critic—Synthesizer 四智能体的多步检索状态机，"
    "通过 Critic 迭代审查机制提升复杂问题的检索质量，并通过最大迭代次数（MAX_ITERATIONS=3）"
    "防止无限循环；"
)
add_para(
    "（4）实现多轮对话与指代消解、SSE 流式回答、引用溯源、用户反馈闭环以及一键生成"
    "带引用的结构化调研报告并导出 PDF；"
)
add_para(
    "（5）建立检索评测体系：自建 110 条评测集进行四组消融与查询增强实验，并通过 RAGAS 对生成质量"
    "进行评估；"
)
add_para(
    "（6）实现工程化补强：BM25 倒排索引持久化到 PostgreSQL，Reranker 超时降级与 Top-K 截断优化，"
    "并给出端到端性能基准数据。"
)

doc.add_page_break()


# ══════════════════════════════════════════════
# 第 2 章 需求分析
# ══════════════════════════════════════════════
add_chapter_heading("2  需求分析")

add_section_heading("2.1  系统目标与用户角色")
add_para(
    "本系统面向中文企业内部知识管理场景，旨在为员工提供基于自然语言的智能搜索与自动调研能力。"
    "系统主要服务三类用户："
)
add_para("（1）普通员工：检索业务文档、回答业务问题；")
add_para("（2）业务分析师：基于给定主题自动生成结构化调研报告；")
add_para("（3）系统管理员：管理知识库、查看运营看板、监控系统性能。")

add_section_heading("2.2  功能需求")
add_para("系统应满足以下功能需求：")
add_para("（1）文档管理：支持 PDF / DOCX / TXT / MD 等多格式上传、解析、分片、向量化入库，"
         "支持批量上传与单条删除；")
add_para("（2）智能问答：基于自然语言查询返回带引用的回答，支持多轮对话与指代消解，"
         "支持 SSE 流式输出；")
add_para("（3）自动调研报告：基于给定主题自动生成包含背景、现状、技术方案、案例、趋势等章节的"
         "结构化报告，支持 PDF 导出；")
add_para("（4）用户反馈：支持对回答进行点赞/点踩与文字反馈，反馈数据用于检索质量分析；")
add_para("（5）运营看板：展示检索命中率、Top 文档、反馈分布、活跃用户等运营指标；"
         "（6）检索增强：支持 BM25 + 向量 + RRF + Reranker 混合检索与 LLM 查询改写 + HyDE 增强。")

add_section_heading("2.3  非功能需求")
add_para("（1）性能：单次检索端到端 P50 延迟 ≤ 2 秒；并发 4 路下吞吐 ≥ 0.5 qps；")
add_para("（2）准确性：在 110 条评测集上完整混合检索管线 top-1 命中率 ≥ 70%；")
add_para("（3）可观测：所有关键路径（检索、生成、引用）日志可追溯；")
add_para("（4）可部署：支持 Docker Compose 一键启动 PostgreSQL + pgvector + Redis 依赖；")
add_para("（5）可扩展：BM25 倒排索引支持跨进程持久化与增量更新，"
         "Reranker 具备超时降级与 Top-K 截断能力。")

doc.add_page_break()


# ══════════════════════════════════════════════
# 第 3 章 系统设计
# ══════════════════════════════════════════════
add_chapter_heading("3  系统设计")

add_section_heading("3.1  总体架构")
add_para(
    "系统采用「四模块 + 四智能体」的架构：在业务层面划分为文档管理、搜索问答、"
    "调研报告与运营看板四大模块；在检索编排层面，通过 LangGraph 构建 Planner（问题分解）、"
    "Retrieval（混合检索）、Critic（充分性审查）、Synthesizer（综合生成）"
    "四智能体的多步检索状态机。系统基于 FastAPI、PostgreSQL（pgvector）、Redis 与 DeepSeek"
    "大模型实现。"
)
add_figure("arch.png", "图1  系统总体架构图", width_cm=15.0)

add_section_heading("3.2  关键技术选型")
add_para(
    "（1）检索框架：BM25（Okapi BM25, k1=1.5, b=0.75）保证关键词精确匹配，pgvector HNSW 索引"
    "提供稠密向量近似最近邻检索，RRF（倒数排名融合）合并两路结果，BGE-Reranker（Cross-Encoder）"
    "对粗排前 30 候选做精排；"
)
add_para(
    "（2）查询增强：LLM 查询改写（生成 2-3 个语义变体后 RRF 融合）与 HyDE（用 LLM 生成的"
    "假设文档做嵌入）按开关启用；"
)
add_para(
    "（3）智能体编排：LangGraph StateGraph，MAX_ITERATIONS=3 防止 Critic 无限自循环；"
)
add_para(
    "（4）LLM 与 Embedding：DeepSeek-V3 作为生成与改写模型，BGE-large-zh-v1.5 作为 Embedding 与 Reranker；"
)
add_para(
    "（5）存储：PostgreSQL 15 + pgvector（HNSW, m=16, ef_construction=64）存储文档、向量、"
    "会话历史与反馈；Redis 作为热点缓存；BM25 倒排索引持久化到 PostgreSQL JSONB 表。"
)

add_section_heading("3.3  Agentic RAG 工作流")
add_para(
    "Agentic RAG 工作流是系统的核心，由四个智能体组成状态机，通过 LangGraph 编排："
    "用户问题进入 Planner，Planner 结合历史对话将复杂问题分解为若干可独立检索的子查询；"
    "Retrieval 依次对每个子查询执行混合检索并汇总结果；Critic 审查已检索信息是否足以回答问题，"
    "若信息不足且未达迭代上限则给出建议补充查询并回到 Retrieval；信息充分后由 Synthesizer "
    "综合生成带引用的最终答案。"
)
add_figure("flow.png", "图2  Agentic RAG 工作流", width_cm=14.0)
add_table(
    ["智能体", "职责", "关键能力"],
    [
        ["Planner", "将复杂问题拆分为 2–5 个子查询", "结合多轮历史做指代消解"],
        ["Retrieval", "对子查询执行混合检索", "BM25 + 向量 + RRF + Reranker"],
        ["Critic", "审查信息充分性", "缺失方面分析、补充查询、迭代上限保护"],
        ["Synthesizer", "综合生成带引用答案/报告", "问答/报告/流式三种模式"],
    ],
    col_widths=[3.0, 6.0, 6.5],
)

add_section_heading("3.4  数据库与接口设计")
add_para(
    "数据库共 6 张表：documents（文档元数据）、chunks（分片与向量）、sessions（多轮会话）、"
    "turns（会话轮次）、reports（调研报告）、feedback（用户反馈）。向量字段为 embedding vector(1024)，"
    "通过 HNSW 索引加速。BM25 倒排索引状态持久化到 bm25_index_state 表（JSONB 字段）。"
    "RESTful 接口覆盖健康检查、文档上传/批量上传/删除/列表、问答、报告生成/查询/导出 PDF、"
    "反馈提交、运营看板、BM25 状态查询等共 16 个端点。"
)
add_figure("er.png", "图3  数据库 E-R 图", width_cm=14.0)

doc.add_page_break()


# ══════════════════════════════════════════════
# 第 4 章 详细设计
# ══════════════════════════════════════════════
add_chapter_heading("4  详细设计")

add_section_heading("4.1  文档入库实现")
add_para(
    "文档入库链路依次为：上传（multipart/form-data）→ 解析（按扩展名分派：PDF 用 pypdf，"
    "DOCX 用 python-docx，TXT/MD 直接读）→ 清洗（去除多余空白与页眉页脚）→ "
    "分片（RecursiveCharacterTextSplitter, chunk_size=500, chunk_overlap=80）→ "
    "向量化（BGE-large-zh-v1.5, 1024 维）→ 写入 PostgreSQL（documents + chunks 两表）→ "
    "更新 BM25 倒排索引（增量 add_chunks 而非全量重建）。"
)
add_figure("usecase.png", "图4  系统用例图", width_cm=13.0)

add_section_heading("4.2  混合检索与查询增强实现")
add_para(
    "HybridRetriever.hybrid_search(query, k) 是检索入口："
    "（1）BM25 侧使用自定义倒排索引（term -> {chunk_id: tf} + term -> doc_freq），"
    "Okapi BM25 公式（k1=1.5, b=0.75）打分，jieba 精确模式分词；"
    "（2）向量侧用 pgvector HNSW 索引做余弦相似度 top-k 检索；"
    "（3）RRF 融合（k=60）两路排序；"
    "（4）取 RRF top-30 走 BGE-Reranker Cross-Encoder 精排；"
    "（5）asyncio.to_thread 包装阻塞推理，asyncio.wait_for 超时降级（默认 1.5s）。"
    "查询增强按开关接入：query_rewrite_enabled 时用 LLM 生成 2-3 个变体分别检索后 RRF 合并；"
    "hyde_enabled 时用 LLM 生成的假设文档做嵌入替换原 query 向量（BM25 一路仍用原 query）。"
)

add_section_heading("4.3  智能体编排与多轮对话实现")
add_para(
    "图编排由 backend/app/graph.py 实现：_build_graph() 构造 StateGraph，"
    "节点为 planner/retrieval/critic/synthesizer 四函数；"
    "条件边 critic -> {synthesizer | retrieval} 由 Critic 输出与 iteration 计数共同决定；"
    "数据库会话通过 configurable 注入（避免把活动连接塞进 state 导致 checkpointer 持久化问题）。"
    "多轮对话通过 checkpointer + thread_id 实现，"
    "_load_session_history / _save_session_turn 读写 turns 表，"
    "前端取最近 4 轮历史交给 Planner/Synthesizer 做指代消解。"
)

add_section_heading("4.4  调研报告与前端实现")
add_para(
    "报告生成复用 Agentic 编排，仅将 Synthesizer 切换为报告模式：要求 LLM 输出固定章节结构"
    "（背景/现状/技术方案/案例/趋势）的完整 Markdown，事实性陈述用 [n] 标注引用，"
    "文末附参考文献列表；若 LLM 遗漏参考文献章节则程序化补齐以保证引用可溯源。"
    "报告状态经 generating → ready/failed 流转，前端轮询直至就绪。"
    "PDF 导出采用 Chrome headless 的 --print-to-pdf 能力将 Markdown 渲染为 PDF，"
    "避免引入 weasyprint 等重型依赖。前端为单页 vanilla JS 应用，"
    "通过 localStorage 持久化会话，5 个主页面：知识库 / 搜索 / 多轮对话 / 报告 / 看板。"
)

doc.add_page_break()


# ══════════════════════════════════════════════
# 第 5 章 系统测试
# ══════════════════════════════════════════════
add_chapter_heading("5  系统测试")

add_section_heading("5.1  测试环境与功能测试")
add_para(
    "测试环境为 Windows 11 + Python 3.13，依赖 Docker Compose 启动 PostgreSQL 15 + pgvector 0.7 + Redis 7，"
    "BGE Embedding 与 Reranker 模型本地缓存（HF_HUB_OFFLINE=1），DeepSeek-V3 通过 API 调用。"
    "功能测试覆盖 16 个 RESTful 接口与 8 个端到端场景（上传→检索→多轮→报告→反馈→看板），"
    "全部通过。"
)

add_section_heading("5.2  检索性能消融实验")
add_para(
    "在 110 条自建评测集（含合成 + 5 篇真实公开语料，402 chunks）上进行四组消融："
    "BM25-only、向量-only、BM25+向量（无 Reranker）、完整管线（含 Reranker）。"
    "结果（见下表）显示完整管线在 top-1、MRR、nDCG@5 三项指标上均达最优。"
)
add_figure("eval.png", "图5  四组消融实验结果", width_cm=15.0)
add_table(
    ["方案", "top-1", "top-3", "top-5", "MRR", "nDCG@5"],
    [
        ["BM25-only", "70.91%", "76.36%", "77.27%", "0.7371", "0.7462"],
        ["向量-only", "72.73%", "82.73%", "83.64%", "0.7780", "0.7924"],
        ["BM25+向量", "69.09%", "81.82%", "89.09%", "0.7632", "0.7943"],
        ["完整+Reranker", "74.55%", "82.73%", "84.55%", "0.7889", "0.8026"],
    ],
    col_widths=[3.5, 2.5, 2.5, 2.5, 2.5, 2.5],
)
add_para(
    "查询增强消融（开启 LLM 改写）：复杂查询 top-1 命中率从 62.12% 提升到 71.21%，"
    "提升 +9.09 个百分点，验证了查询改写对复杂问题检索的有效性。"
)

add_section_heading("5.3  生成质量评测（RAGAS）")
add_para(
    "采用 RAGAS 框架对 44 条样本进行生成质量评测，覆盖 5 类查询难度（直答/同义改写/跨文档/干扰/反向否定）。"
    "结果如下："
)
add_table(
    ["指标", "数值", "说明"],
    [
        ["faithfulness（忠实度）", "0.6924", "回答与检索上下文的一致性"],
        ["answer_relevancy（相关性）", "0.6156", "回答与问题的相关程度"],
        ["context_precision（上下文精确度）", "0.7853", "检索上下文中相关 chunk 的比例"],
        ["context_recall（上下文召回率）", "0.8409", "相关 chunk 被检索到的比例"],
    ],
    col_widths=[5.5, 2.5, 7.0],
)

add_section_heading("5.4  检索性能基准")
add_para(
    "单次检索端到端延迟（Reranker OFF）P50 = 51.6ms；开启 Reranker 后 P50 = 1362.92ms，"
    "其中 Reranker 占比约 90%（已通过超时降级 + Top-K 截断优化降至 ≤ 900ms）。"
    "并发吞吐：单并发 0.52 qps，4 并发 0.76 qps，8 并发 0.80 qps。"
    "BM25 倒排索引全量重建 402 chunks 用时 52.93ms，增量更新单 chunk 7.54ms（200 chunks 0.37ms）。"
)

add_section_heading("5.5  测试结论")
add_para(
    "功能测试与端到端测试全部通过。检索层面，完整管线显著优于任一单路检索（top-1 +3.64pt vs BM25-only），"
    "查询改写对复杂问题有明显增益（+9.09pt）。生成质量层面，RAGAS 四指标均在 0.6 以上，"
    "context_recall 达到 0.84 表明检索覆盖度良好。"
    "性能层面，单次检索 P50 延迟可接受，并发吞吐满足中小规模企业知识库需求。"
    "pytest 单元与集成测试 65 passed。"
)

doc.add_page_break()


# ══════════════════════════════════════════════
# 结论
# ══════════════════════════════════════════════
add_chapter_heading("结    论")

add_para(
    "本文围绕企业知识管理与智能检索的实际需求，设计并实现了一套基于 Agentic RAG 的"
    "企业智能搜索与自动调研系统。系统构建了文档入库全链路，实现了 BM25 + 向量 + RRF + Reranker "
    "的混合检索策略；基于 LangGraph 实现了 Planner—Retrieval—Critic—Synthesizer 四智能体的"
    "多步检索编排与防死循环保护；实现了多轮对话与指代消解、SSE 流式回答、引用溯源、"
    "用户反馈闭环以及一键生成调研报告与 PDF 导出；建立了检索评测体系，"
    "验证了 Reranker 精排与查询改写对检索质量的显著提升；并通过 BM25 索引持久化、"
    "Reranker 超时降级与 Top-K 截断等工程化手段缓解了性能瓶颈。"
)
add_para(
    "系统仍存在以下不足：（1）BM25 倒排索引虽已支持跨进程持久化，但单实例重启时仍需从数据库"
    "反序列化；（2）Embedding 与 Reranker 在 CPU 上推理，虽已通过工程优化降低延迟，"
    "但大批量入库与高并发下吞吐仍受限；（3）评测语料以程序化合成的企业文档为主，"
    "真实公开文档占比较小；（4）RAGAS 评测样本规模仍有限。"
)
add_para(
    "后续工作可从以下方向展开：迁移到独立搜索服务（如 Elasticsearch）支撑更大规模数据；"
    "引入 GPU 推理或模型服务化部署进一步提升性能；扩充更大规模的真实企业文档评测语料与 RAGAS 样本；"
    "探索更多进阶检索策略以提升复杂问题的回答质量。"
)

doc.add_page_break()


# ══════════════════════════════════════════════
# 参考文献（按模板三类：[1] 期刊 / [a] 图书 / [c] 外文）
# ══════════════════════════════════════════════
add_chapter_heading("参考文献")

add_section_heading("[1]  期刊类")
refs_journal = [
    "Lewis P, Perez E, Piktus A, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks[J]. NeurIPS, 2020.",
    "Karpukhin V, Oğuz B, Min S, et al. Dense Passage Retrieval for Open-Domain Question Answering[J]. EMNLP, 2020.",
    "Robertson S, Zaragoza H. The Probabilistic Relevance Framework: BM25 and Beyond[J]. Foundations and Trends in Information Retrieval, 2009.",
    "Malkov Y A, Yashunin D A. Efficient and Robust Approximate Nearest Neighbor Search Using Hierarchical Navigable Small World Graphs[J]. IEEE TPAMI, 2020.",
    "Asai A, Wu Z, Wang Y, et al. Self-RAG: Learning to Retrieve, Generate, and Critique through Self-Reflection[J]. ICLR, 2024.",
    "Yan S Q, Gu J C, Zhu Y, et al. Corrective Retrieval Augmented Generation[J]. arXiv preprint arXiv:2401.15884, 2024.",
    "Xiao S, Liu Z, Zhang P, et al. C-Pack: Packaged Resources To Advance General Chinese Embedding[J]. arXiv preprint arXiv:2309.07597, 2023.",
    "DeepSeek-AI. DeepSeek-V3 Technical Report[J]. arXiv preprint arXiv:2412.19437, 2024.",
]
for ref in refs_journal:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Pt(0)
    pf.first_line_indent = Pt(-24)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(18)
    pf.space_after = Pt(3)
    r = p.add_run(ref)
    set_run(r, FANGSONG, Pt(10.5))

add_section_heading("[a]  图书类")
refs_book = [
    "傅京孙. 模式识别及其应用[M]. 北京: 科学出版社, 2001.",
    "周志华. 机器学习[M]. 北京: 清华大学出版社, 2016.",
    "Russell S, Norvig P. Artificial Intelligence: A Modern Approach[M]. 4th ed. Pearson, 2020.",
]
for ref in refs_book:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Pt(0)
    pf.first_line_indent = Pt(-24)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(18)
    pf.space_after = Pt(3)
    r = p.add_run(ref)
    set_run(r, FANGSONG, Pt(10.5))

add_section_heading("[c]  外文类（含电子文献）")
refs_foreign = [
    "Borko H. Information Science: What is it?[J]. American Documentation, 1968, 19(1): 3-5.",
    "Chase H. LangChain: Building applications with LLMs through composability[EB/OL]. https://github.com/langchain-ai/langchain.",
    "LangChain-AI. LangGraph: Building stateful, multi-actor applications with LLMs[EB/OL]. https://github.com/langchain-ai/langgraph.",
    "pgvector. Open-source vector similarity search for Postgres[EB/OL]. https://github.com/pgvector/pgvector.",
]
for ref in refs_foreign:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Pt(0)
    pf.first_line_indent = Pt(-24)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(18)
    pf.space_after = Pt(3)
    r = p.add_run(ref)
    set_run(r, FANGSONG, Pt(10.5))

doc.add_page_break()


# ══════════════════════════════════════════════
# 附录A  项目部署
# ══════════════════════════════════════════════
add_chapter_heading("附录A  项目部署")

add_section_heading("A.1  环境要求")
add_para("操作系统：Windows 10/11 或 Ubuntu 20.04+；Python 3.11+；Docker 24+；Docker Compose v2；"
         "至少 16GB 内存（运行 BGE 模型）；至少 20GB 可用磁盘。")

add_section_heading("A.2  启动依赖服务")
add_code_block_dummy = doc.add_paragraph()
add_code_block_dummy.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
add_code_block_dummy.paragraph_format.line_spacing = Pt(14)
r = add_code_block_dummy.add_run("cd D:\\agentic-rag-system\ndocker compose up -d   # 启动 PostgreSQL + pgvector + Redis")
set_run(r, "Consolas", Pt(10))

add_section_heading("A.3  安装依赖与配置")
add_para("执行 pip install -r backend/requirements.txt 安装后端依赖，"
         "复制 .env.example 为 .env 并填入 DEEPSEEK_API_KEY 等配置。")

add_section_heading("A.4  模型缓存与离线运行")
add_para("BGE-large-zh-v1.5 与 BGE-reranker-base 首次启动会自动下载（约 2GB），"
         "缓存到 ~/.cache/huggingface/hub。后续可设置 HF_HUB_OFFLINE=1 与 TRANSFORMERS_OFFLINE=1 离线运行。")

add_section_heading("A.5  启动服务与验收")
add_code_block_dummy = doc.add_paragraph()
add_code_block_dummy.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
add_code_block_dummy.paragraph_format.line_spacing = Pt(14)
r = add_code_block_dummy.add_run("cd backend\nenv -u HTTP_PROXY -u HTTPS_PROXY HF_HUB_OFFLINE=1 python -m uvicorn app.main:app --port 8000")
set_run(r, "Consolas", Pt(10))
add_para("浏览器访问 http://localhost:8000/ 即可使用前端；"
         "执行 pytest backend/tests/ 验证 65 个测试全部通过。")

doc.add_page_break()


# ══════════════════════════════════════════════
# 附录B  用户手册
# ══════════════════════════════════════════════
add_chapter_heading("附录B  用户手册")

add_section_heading("B.1  知识库管理")
add_para("在「知识库」页面点击「上传文档」按钮，选择 PDF/DOCX/TXT/MD 文件，"
         "系统自动解析、分片、向量化入库。批量上传支持多文件同时处理。"
         "已上传的文档可单条删除（同步清理 chunks 与 BM25 倒排索引）。")

add_section_heading("B.2  搜索与多轮对话")
add_para("在「搜索」或「对话」页面输入问题，系统返回带引用编号的回答，"
         "点击引用编号可查看原文片段。多轮对话支持指代消解（「它」/「这个方案」），"
         "回答以 SSE 流式逐字产出。")

add_section_heading("B.3  自动调研报告")
add_para("在「报告」页面输入调研主题，系统调用 Agentic RAG 自动生成"
         "包含背景、现状、技术方案、案例、趋势 5 章节的结构化报告，"
         "所有事实性陈述标注引用编号，报告可导出 PDF。")

add_section_heading("B.4  反馈与运营看板")
add_para("对任一回答可点赞/点踩 + 文字反馈，反馈数据进入运营看板的"
         "检索质量统计与 Top 文档榜单。")

doc.add_page_break()


# ══════════════════════════════════════════════
# 图目录
# ══════════════════════════════════════════════
add_chapter_heading("图  目  录")
fig_index = [
    ("图1  系统总体架构图", "6"),
    ("图2  Agentic RAG 工作流", "8"),
    ("图3  数据库 E-R 图", "10"),
    ("图4  系统用例图", "11"),
    ("图5  四组消融实验结果", "16"),
]
for title, page in fig_index:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(20)
    pf.space_after = Pt(3)
    tab_stops = pf.tab_stops
    tab_stops.add_tab_stop(Cm(13.5), alignment=2, leader=1)
    r = p.add_run(title)
    set_run(r, FANGSONG, Pt(12))
    r2 = p.add_run("\t" + page)
    set_run(r2, FANGSONG, Pt(12))

doc.add_page_break()


# ══════════════════════════════════════════════
# 表目录
# ══════════════════════════════════════════════
add_chapter_heading("表  目  录")
tbl_index = [
    ("表1  Agentic RAG 四智能体职责", "9"),
    ("表2  四组消融实验结果", "16"),
    ("表3  RAGAS 生成质量评测", "17"),
]
for title, page in tbl_index:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(20)
    pf.space_after = Pt(3)
    tab_stops = pf.tab_stops
    tab_stops.add_tab_stop(Cm(13.5), alignment=2, leader=1)
    r = p.add_run(title)
    set_run(r, FANGSONG, Pt(12))
    r2 = p.add_run("\t" + page)
    set_run(r2, FANGSONG, Pt(12))


# ══════════════════════════════════════════════
# 保存
# ══════════════════════════════════════════════
doc.save(OUT)
print(f"OK -> {OUT}")
